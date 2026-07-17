#!/usr/bin/env python3
"""Convert input documents to Markdown for downstream processing.

Supports: txt, md, docx
Output: input_md/<basename>.md

Usage: python3 scripts/convert_to_md.py input/
"""

import os
import sys


def docx_to_md(filepath):
    """Convert DOCX to Markdown, preserving headings, bold, italic, lists."""
    from docx import Document
    from docx.oxml.ns import qn
    import re

    doc = Document(filepath)
    lines = []

    def _get_formatted_line(para):
        """Extract paragraph text with bold/italic inline formatting, preserving line breaks."""
        result = []
        for run in para.runs:
            # Check if this run contains a br element (line break)
            brs = run._element.findall(qn('w:br'))
            if brs:
                # Add newlines for each br
                result.append("\n" * len(brs))
            t = run.text
            if not t:
                continue
            if run.bold and run.italic:
                t = f"***{t}***"
            elif run.bold:
                t = f"**{t}**"
            elif run.italic:
                t = f"*{t}*"
            result.append(t)
        return "".join(result) if result else para.text

    def _is_list_paragraph(para):
        """Check if paragraph belongs to a numbered/bulleted list."""
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                return True
        return False

    def _get_list_level(para):
        """Get list nesting level (0-based)."""
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                ilvl = numPr.find(qn('w:ilvl'))
                if ilvl is not None:
                    return int(ilvl.get(qn('w:val'), 0))
        return 0

    def _detect_heading_by_style(para):
        """Return heading level (1-6) based on paragraph style, or None."""
        if para.style is None:
            return None
        style_name = (para.style.name or "").lower()
        if style_name.startswith("heading 1") or style_name.startswith("heading1"):
            return 1
        elif style_name.startswith("heading 2") or style_name.startswith("heading2"):
            return 2
        elif style_name.startswith("heading 3") or style_name.startswith("heading3"):
            return 3
        elif style_name.startswith("heading 4") or style_name.startswith("heading4"):
            return 4
        elif style_name.startswith("heading"):
            return 5
        # Check outline level (lvl=0 → heading 1)
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            outlineLvl = pPr.find(qn('w:outlineLvl'))
            if outlineLvl is not None:
                level = int(outlineLvl.get(qn('w:val'), 0)) + 1
                return min(level, 6)
        return None

    def _detect_heading_by_content(text):
        """Return heading level based on Chinese content patterns, or None.

        Detects patterns like:
          - 一、Title  /  二、Title  (→ level 2)
          - 第一章 xxx  /  第二节 xxx  (→ level 1)
          - 1. Title  /  1、Title  (→ level 2)
          - (一) Title  /  （二）Title  (→ level 3)
        """
        if not text:
            return None
        # 1. Title
        if re.match(r'^#\s+\S', text):
            return 1
        # 第X章 / 第X节 / 第X篇 / 第X部  (top-level)
        if re.match(r'^第[一二三四五六七八九十百千\d]+[章节篇部课讲]\s', text):
            return 1
        # 一、 / 二、 / 三、 (Chinese numbered)
        if re.match(r'^[一二三四五六七八九十]+[、，,]\.?\s*\S', text):
            return 2
        # 1. / 1、/ 1)  (Arabic numbered)
        if re.match(r'^\d+[\.、\)]\s*\S', text):
            return 2
        # (一) / （一）/ (1) / ①
        if re.match(r'^[（(\[][一二三四五六七八九十\d]+[）)\]][\s\S]', text):
            return 3
        # 1.1 / 1.1.1  (dotted)
        if re.match(r'^\d+\.\d+', text):
            return 3
        # Short bold-only line (likely a heading without numbering)
        return None

    def _detect_heading_by_font(para):
        """Return heading level based on font size and boldness, or None.

        Detects headings that use visual formatting (large/bold text)
        instead of proper heading styles. Only triggers when the
        paragraph is relatively short (typical heading length).
        """
        max_size_pt = None
        all_bold = True
        has_text = False

        for run in para.runs:
            t = (run.text or "").strip()
            if not t:
                continue
            has_text = True
            if not run.bold:
                all_bold = False
            if run.font.size is not None:
                # python-docx stores size in EMU; 1 pt = 12700 EMU
                size_pt = run.font.size / 12700
                if max_size_pt is None or size_pt > max_size_pt:
                    max_size_pt = size_pt

        # Fallback: check paragraph-level font size (w:pPr/w:rPr/w:sz)
        # Many DOCX files define font size in the paragraph style, not on each run
        if max_size_pt is None:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                rPr = pPr.find(qn('w:rPr'))
                if rPr is not None:
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        # w:sz value is in half-points (e.g., 48 = 24pt)
                        max_size_pt = int(sz.get(qn('w:val'))) / 2

        if not has_text or max_size_pt is None:
            return None

        text = para.text.strip()
        # Headings are typically short; skip long paragraphs
        if len(text) > 80:
            return None

        # Map font sizes to heading levels
        if max_size_pt >= 22:
            return 1
        elif max_size_pt >= 18:
            return 2
        elif max_size_pt >= 15:
            return 3
        elif max_size_pt >= 13 and all_bold:
            return 4

        return None

    for para in doc.paragraphs:
        # Get formatted text with inline formatting and line breaks
        line = _get_formatted_line(para)
        text = line.strip()
        if not text:
            lines.append("")
            continue

        is_list = _is_list_paragraph(para)

        # Try style-based heading detection first
        heading_level = _detect_heading_by_style(para)
        # Fall back to content-based detection
        if heading_level is None:
            heading_level = _detect_heading_by_content(text)
        # Fall back to font-size-based detection
        if heading_level is None:
            heading_level = _detect_heading_by_font(para)

        if heading_level:
            hashes = "#" * heading_level
            # Normalize internal line breaks (w:br in docx runs) to spaces
            # so the heading stays on one line. Otherwise AI can't detect the
            # heading pattern during split phase.
            clean_text = text.replace("\n", " ").replace("\r", "")
            lines.append(f"{hashes} {clean_text}")
        elif is_list:
            list_level = _get_list_level(para)
            indent = "  " * list_level
            lines.append(f"{indent}- {text}")
        else:
            lines.append(line)

    # Tables
    for table in doc.tables:
        lines.append("")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            # Header separator after first row
            if row == table.rows[0]:
                lines.append("|" + "|".join(["---"] * len(cells)) + "|")
        lines.append("")
    return "\n".join(lines)


CONVERTERS = {
    ".docx": docx_to_md,
}


def convert_file(filepath, output_dir):
    """Convert a single file to .md in output_dir. Returns output path."""
    ext = os.path.splitext(filepath)[1].lower()
    basename = os.path.splitext(os.path.basename(filepath))[0]
    output_path = os.path.join(output_dir, basename + ".md")

    if ext in (".md", ".txt"):
        # Already plain text; copy
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        print(f"  copy: {filepath} -> {output_path} ({len(content)} chars)")
        return output_path

    converter = CONVERTERS.get(ext)
    if converter is None:
        print(f"  skip: {filepath} (unsupported format {ext})")
        return None

    try:
        content = converter(filepath)
    except ModuleNotFoundError as e:
        print(f"  ERROR converting {filepath}: {e}")
        print(f"  HINT: Required Python package is not installed in this environment.")
        print(f"  DO NOT attempt pip install or write helper scripts — report this error and skip the file.")
        return None
    except Exception as e:
        print(f"  ERROR converting {filepath}: {e}")
        return None

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"  convert: {filepath} -> {output_path} ({len(content)} chars)")
    return output_path


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/convert_to_md.py <input_dir>")
        sys.exit(1)

    input_dir = sys.argv[1]
    if not os.path.isdir(input_dir):
        print(f"Error: {input_dir} is not a directory")
        sys.exit(1)

    output_dir = "input_md"
    os.makedirs(output_dir, exist_ok=True)

    # ── Discover input files ──
    input_files = sorted(
        f for f in os.listdir(input_dir)
        if os.path.isfile(os.path.join(input_dir, f))
        and not f.startswith(".")
    )

    print(f"[convert_to_md] input_dir = {input_dir}")
    print(f"[convert_to_md] output_dir = {output_dir}/")
    print(f"[convert_to_md] found {len(input_files)} file(s):")
    for f in input_files:
        size = os.path.getsize(os.path.join(input_dir, f))
        print(f"  - {f} ({size:,} bytes)")

    if not input_files:
        print("[convert_to_md] SKIP: no files to process")
        return

    # ── Convert ──
    converted = []
    skipped = []
    errors = []

    for name in input_files:
        filepath = os.path.join(input_dir, name)
        result = convert_file(filepath, output_dir)
        if result:
            converted.append(result)
        else:
            ext = os.path.splitext(name)[1].lower()
            if ext not in CONVERTERS and ext not in (".md", ".txt"):
                skipped.append(name)
            else:
                errors.append(name)

    # ── Summary ──
    print()
    print(f"[convert_to_md] SUCCESS: {len(converted)} converted, {len(skipped)} skipped, {len(errors)} errors")

    if converted:
        print(f"[convert_to_md] output files in {output_dir}/:")
        for p in converted:
            size = os.path.getsize(p)
            basename = os.path.basename(p)
            print(f"  -> {basename} ({size:,} bytes)")

    if skipped:
        print(f"[convert_to_md] skipped (unsupported format):")
        for f in skipped:
            print(f"  - {f}")

    if errors:
        print(f"[convert_to_md] errors:")
        for f in errors:
            print(f"  - {f}")


if __name__ == "__main__":
    main()
